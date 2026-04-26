from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "心理支持Agent训练集收集手册.docx"

BLUE = RGBColor(23, 54, 93)
MID_BLUE = RGBColor(47, 85, 151)
GRAY = RGBColor(91, 103, 122)


def set_run(run, *, size: int = 10, bold: bool = False, color: RGBColor | None = None, font: str = "Microsoft YaHei") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("心理支持 Agent 训练集收集手册")
    set_run(run, size=22, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("适用于校园心理支持系统：多轮对话、心理熵、风险评估、支持方案与转介判断")
    set_run(run, size=12, color=GRAY)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("版本：v1.0    用途：组内数据采集、标注规范、后续 SFT/评估数据准备")
    set_run(run, size=10)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run(run, size=16 if level == 1 else 12, bold=True, color=BLUE if level == 1 else MID_BLUE)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run(run, size=10)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.6)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(f"• {item}")
        set_run(run, size=10)


def add_card(doc: Document, title: str, fields: list[tuple[str, str]]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"【{title}】")
    set_run(run, size=11, bold=True, color=MID_BLUE)
    for label, value in fields:
        line = doc.add_paragraph()
        line.paragraph_format.left_indent = Cm(0.5)
        line.paragraph_format.space_after = Pt(1)
        r1 = line.add_run(f"{label}：")
        set_run(r1, size=9, bold=True, color=BLUE)
        r2 = line.add_run(value)
        set_run(r2, size=9)


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run(run, size=8, font="Consolas")


def build_docx() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    add_title(doc)

    add_heading(doc, "一、收集目标")
    add_body(
        doc,
        "本手册用于指导项目组收集高质量训练集。当前阶段不要盲目追求数据量，重点是让模型先具备自然多轮对话能力，并逐步沉淀项目专属的风险、心理熵、策略与转介标签。",
    )
    add_bullets(
        doc,
        [
            "让本地模型能自然接话，不再复读模板或把弱输入当选择题。",
            "覆盖校园常见心理压力场景，例如考试、宿舍、人际、隐私、睡眠、未来焦虑。",
            "为后端心理熵、风险分层、转介触发和动态跟踪提供可解释数据。",
            "为后续 SFT、评估集、偏好训练或 KTO/DPO 准备干净样本。",
        ],
    )

    add_heading(doc, "二、训练集总览")
    overview = [
        ("多轮自然聊天", "让模型正常接话，不机械复读", "是", "最高", "5k-20k 轮", "messages 多轮对话"),
        ("校园心理支持对话", "贴合校园心理陪伴助手", "是", "最高", "3k-10k 轮", "user + assistant"),
        ("弱输入承接", "解决“嗯/哦/1/？”不会接", "是", "最高", "1k-3k 条", "单轮或多轮"),
        ("场景化压力数据", "覆盖考试、宿舍、人际、家庭、未来焦虑", "是", "最高", "5k-20k 条", "多轮对话"),
        ("隐私与信任数据", "处理“不想说、怕别人知道”", "是", "高", "1k-3k 条", "多轮对话"),
        ("心理熵标注数据", "验证熵值系统和维度分数", "是", "高", "1k-5k 条", "输入 + 熵标签"),
        ("风险等级标注数据", "判断 low/medium/high/critical", "是", "高", "1k-5k 条", "输入 + 风险等级"),
        ("支持方案数据", "输出具体建议而不是空话", "是", "高", "2k-8k 条", "问题 + 支持方案"),
        ("转介判断数据", "判断何时建议老师/心理中心/紧急求助", "是", "高", "500-2k 条", "输入 + 是否转介"),
        ("高风险危机场景", "自伤、自杀、伤人、极端绝望安全处理", "谨慎收集", "高", "300-1k 条", "风险输入 + 安全回复"),
        ("日常闲聊", "避免一上来就心理分析", "建议", "中", "3k-10k 条", "多轮聊天"),
        ("拒绝诊断数据", "避免模型乱下诊断", "建议", "中", "500-2k 条", "用户问诊断 + 安全回复"),
        ("坏回复对比数据", "后期 DPO/KTO 使用", "暂缓", "中后期", "500-3k 对", "chosen/rejected"),
    ]
    for name, purpose, required, priority, amount, form in overview:
        add_card(
            doc,
            name,
            [
                ("主要目的", purpose),
                ("是否必须", required),
                ("优先级", priority),
                ("建议数量", amount),
                ("数据形式", form),
            ],
        )

    add_heading(doc, "三、八类核心数据说明")
    core = [
        ("多轮自然聊天", "日常、兴趣、学习、闲聊转心理话题", "3-8 轮以上，能自然承接上下文；不要每轮都像心理咨询。"),
        ("校园心理支持", "考试焦虑、宿舍矛盾、社交孤立、家庭压力、未来迷茫", "用户口语化，助手温和、具体、不过度说教。"),
        ("弱输入承接", "嗯、哦、？、1、不知道、不想说、随便、算了", "不能复读、不能加数字、不能强行追问，要结合上下文。"),
        ("隐私与信任", "怕别人知道、怕告诉老师、不想说具体是谁", "先解释安全边界，再降低表达压力。"),
        ("心理熵标注", "情绪强度、认知负荷、睡眠饮食、人际张力等", "每条包含 entropy_score、level、balance_state、dominant_drivers。"),
        ("风险等级", "low/medium/high/critical", "low 和 medium 也很重要，不能只收危机场景。"),
        ("支持方案", "总结、当下支持、自我调节、校园资源、追问", "建议必须具体、低成本、可执行。"),
        ("转介判断", "是否建议心理中心、辅导员、校医院、紧急求助", "包含 should_refer、urgency、recommended_channel、reasons。"),
    ]
    for title, scene, standard in core:
        add_card(doc, title, [("覆盖场景", scene), ("质量要求", standard)])

    add_heading(doc, "四、推荐标注字段模板")
    fields = [
        ("id", "样本唯一编号，例如 campus_000001。"),
        ("scenario", "场景类别，例如考试焦虑、宿舍冲突、隐私担心。"),
        ("messages", "多轮对话内容，使用 system/user/assistant 格式。"),
        ("visible_reply", "用户实际看到的回复，应自然、支持性、不过度专业化。"),
        ("risk_level", "风险等级：low / medium / high / critical。"),
        ("entropy_score", "心理熵分数，范围 0-100。"),
        ("entropy_level", "心理熵等级，范围 1-5。"),
        ("balance_state", "状态：stable / strained / fragile / crisis。"),
        ("dominant_drivers", "主要熵驱动，例如认知负荷、睡眠饮食、人际张力。"),
        ("support_strategy", "支持策略，例如安抚、澄清、拆任务、转介。"),
        ("should_refer", "是否建议转介：true / false。"),
        ("urgency", "转介紧急度：none / watch / recommended / urgent。"),
        ("quality_score", "人工质量评分，建议 1-5。"),
        ("notes", "标注备注，例如是否自然、是否适合 SFT。"),
    ]
    for key, desc in fields:
        add_card(doc, key, [("说明", desc)])

    add_heading(doc, "五、Excel 采集样例")
    examples = [
        ("001", "考试焦虑", "我怕挂科", "最近考试很多", "听起来你最近一直被考试压着，我们先把复习范围缩小一点。", "medium", "45", "认知负荷", "否", "none"),
        ("002", "隐私担心", "我怕别人知道", "无", "你放心，我们先停在你觉得安全的范围里，不用说具体名字。", "low", "30", "信任安全感", "否", "none"),
        ("003", "睡眠饮食", "睡不好也吃不下", "最近压力大", "这已经影响到睡眠和吃饭了，我们先从一个最小动作开始稳住。", "medium", "62", "生理失衡", "是", "recommended"),
        ("004", "危机风险", "我不想活了", "多轮负面表达", "我很担心你的安全，现在请先联系身边可信任的人或紧急支持。", "critical", "90", "风险压力", "是", "urgent"),
    ]
    for item in examples:
        add_card(
            doc,
            f"样例 {item[0]} - {item[1]}",
            [
                ("用户输入", item[2]),
                ("历史上下文", item[3]),
                ("好回复", item[4]),
                ("风险等级", item[5]),
                ("心理熵分数", item[6]),
                ("主要驱动", item[7]),
                ("是否转介", item[8]),
                ("紧急度", item[9]),
            ],
        )

    add_heading(doc, "六、示例 JSON 格式")
    add_code(
        doc,
        r'''
{
  "messages": [
    {"role": "user", "content": "我最近压力很大，晚上睡不好。"},
    {"role": "assistant", "content": "听起来这段时间你一直绷着，连睡觉都放松不下来。"},
    {"role": "user", "content": "嗯"},
    {"role": "assistant", "content": "那我们先不急着讲很多。你只要告诉我，现在更明显的是睡不着，还是脑子停不下来？"}
  ],
  "risk_level": "medium",
  "entropy_score": 52,
  "entropy_level": 3,
  "balance_state": "strained",
  "dominant_drivers": ["生理失衡(睡不好)", "认知负荷(压力)"],
  "should_refer": false,
  "urgency": "none"
}
''',
    )

    add_heading(doc, "七、好数据与坏数据标准")
    quality = [
        ("多轮、有上下文", "单句问答太多"),
        ("用户表达像真实学生", "全是书面语、模板化"),
        ("回复自然、有人味", "回复像心理学论文或分析报告"),
        ("有具体下一步", "只会说加油、放松、别想太多"),
        ("尊重边界", "一直追问细节"),
        ("能处理弱输入", "遇到“1”就回答“2”"),
        ("不乱诊断", "随便说你抑郁了"),
        ("高风险能安全处理", "自杀表达还普通闲聊"),
    ]
    for good, bad in quality:
        add_card(doc, "质量对照", [("好数据", good), ("坏数据", bad)])

    add_heading(doc, "八、第一阶段建议收集量")
    first_stage = [
        ("多轮自然聊天", "3000 条", "先解决模型不会聊天的问题"),
        ("校园心理支持", "3000 条", "考试、宿舍、人际、家庭、未来焦虑"),
        ("弱输入承接", "1000 条", "重点解决嗯、哦、？、1、不想说"),
        ("隐私信任场景", "500 条", "怕别人知道、怕被记录、不想说具体细节"),
        ("睡眠/考试/宿舍/人际", "3000 条", "最贴合校园心理项目"),
        ("风险等级标注", "1000 条", "覆盖 low 到 critical"),
        ("心理熵标注", "1000 条", "用于验证和改进熵值系统"),
        ("转介判断", "500 条", "should_refer 和 urgency"),
    ]
    for title, amount, note in first_stage:
        add_card(doc, title, [("建议数量", amount), ("备注", note)])

    add_heading(doc, "九、组员分工建议")
    add_bullets(
        doc,
        [
            "A 组：负责多轮自然聊天和日常闲聊，重点检查是否自然。",
            "B 组：负责校园心理场景，包括考试、宿舍、人际、家庭压力。",
            "C 组：负责弱输入、隐私担心、不想说等低表达样本。",
            "D 组：负责风险等级、心理熵、转介字段标注。",
            "每周抽样复查 100 条，删除机械安慰、过度分析、乱诊断、无具体帮助的样本。",
        ],
    )

    add_heading(doc, "十、当前阶段结论")
    add_body(
        doc,
        "当前最重要的不是继续下载大量公开数据，而是积累项目自己的高质量结构化样本。公开数据可以辅助，但真正体现项目创新的是“对话 + 心理熵 + 风险 + 策略 + 转介”的闭环数据。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_docx()
    print(OUT)
